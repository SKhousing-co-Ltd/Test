"""Build a Document Generation template that keeps the approved restoration tables.

The screen still stores the editable restoration criteria as a contract-specific
text value.  The approved, standard construction/equipment criteria are laid
out in two Word tables; keeping them as template content preserves their
print layout in both the generated DOCX and formal PDF.  The Edge Function
merges only contract-specific additions into ``{{restorationNotes}}``.
"""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


def insert_after(reference, element) -> None:
    reference.addnext(element)


def paragraph_with_text(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def build(tagged_template: Path, approved_source: Path, destination: Path) -> None:
    target = Document(tagged_template)
    source = Document(approved_source)

    restoration_tag = paragraph_with_text(target, "{{restorationText}}")
    restoration_tag.clear()
    restoration_tag.add_run("{{restorationNotes}}")

    construction_heading = paragraph_with_text(source, "■ 建築")
    equipment_heading = paragraph_with_text(source, "■ 設備")
    if len(source.tables) < 3:
        raise ValueError("The approved source must contain contract, construction, and equipment tables.")

    # Clone WordprocessingML rather than recreating tables. This keeps the
    # approved column widths, borders, merged cells, fonts, and row wrapping.
    anchor = restoration_tag._element
    for element in (
        deepcopy(construction_heading._element),
        deepcopy(source.tables[1]._element),
        deepcopy(equipment_heading._element),
        deepcopy(source.tables[2]._element),
    ):
        insert_after(anchor, element)
        anchor = element

    destination.parent.mkdir(parents=True, exist_ok=True)
    target.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("tagged_template", type=Path)
    parser.add_argument("approved_source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()
    build(args.tagged_template, args.approved_source, args.destination)


if __name__ == "__main__":
    main()
