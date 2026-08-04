"""Create the ordinary-lease Adobe Document Generation demo template.

The legal team should replace this with the approved DOCX, preserving the
template tags listed below.  The script keeps the demo template reproducible.
"""
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Mm, Pt
from docx.oxml.ns import qn

OUTPUT = Path("tmp/ordinary_lease_document_generation_v1.docx")


def set_font(run, size=10.5, bold=False):
    run.font.name = "Yu Mincho"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Mincho")
    run.font.size = Pt(size)
    run.font.bold = bold


def paragraph(document, text="", size=10.5, bold=False):
    item = document.add_paragraph()
    item.paragraph_format.space_after = Pt(4)
    item.paragraph_format.line_spacing = 1.35
    set_font(item.add_run(text), size=size, bold=bold)
    return item


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = section.bottom_margin = Mm(20)
    section.left_margin = section.right_margin = Mm(18)

    heading = paragraph(document, "建物賃貸借契約書", size=16, bold=True)
    heading.alignment = 1
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, tag in [
        ("賃借人", "{{tenantName}}"), ("物件名", "{{propertyName}}"),
        ("所在地", "{{propertyAddress}}"), ("貸室・区画", "{{unitNames}}"),
        ("契約期間", "{{contractStartDate}} から {{contractEndDate}}"),
        ("賃料", "{{monthlyRentAmount}}"), ("使用目的", "{{usePurpose}}"),
        ("特約", "{{specialProvisions}}"),
    ]:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = label, tag
        for cell in cells:
            for run in cell.paragraphs[0].runs:
                set_font(run)

    paragraph(document, "契約約款", size=13, bold=True)
    paragraph(document, "{{termsText}}")
    paragraph(document, "原状回復工事基準（別表）", size=13, bold=True)
    paragraph(document, "{{restorationText}}")
    paragraph(document, "署名・押印", size=13, bold=True)
    paragraph(document, "賃貸人　____________________________　　賃借人　____________________________")
    paragraph(document, "対象箇所図", size=13, bold=True)
    paragraph(document, "※ 正式ひな型では Adobe Document Generation Word Add-in で planImage の画像タグを配置します。")
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
