"""Create the ordinary-lease Adobe Document Generation DOCX template.

The script preserves the approved Word layout and replaces only contract-specific
content with Document Generation tags.  The contract terms and restoration
criteria are deliberately rendered as standalone paragraphs so long Japanese
text can flow over pages without Acrobat form-field limits.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.table import _Cell
from docx.text.paragraph import Paragraph


def replace_paragraph(paragraph: Paragraph, value: str) -> None:
    """Replace paragraph contents but retain its paragraph-level formatting."""
    paragraph.clear()
    paragraph.add_run(value)


def replace_cell(cell: _Cell, value: str) -> None:
    """Replace a table cell while keeping its borders, alignment and shading."""
    replace_paragraph(cell.paragraphs[0], value)
    for paragraph in cell.paragraphs[1:]:
        replace_paragraph(paragraph, "")


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def remove_table(table) -> None:
    table._element.getparent().remove(table._element)


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def create_template(source: Path, destination: Path) -> None:
    document = Document(source)
    terms_table = document.tables[0]

    cell_values = {
        2: "{{tenantName}}",
        3: "{{guarantorName}}（極度額    円を上限とする）",
        4: "{{propertyName}}",
        5: "{{propertyLotAddress}}",
        6: "{{propertyAddress}}",
        7: "{{buildingStructure}}",
        8: "{{floorLabel}}階    {{leasedAreaSqm}}㎡（{{unitNames}}）",
        9: "{{usePurpose}}",
        10: "{{contractStartDate}}から    {{contractEndDate}}まで",
        11: "月額金    {{monthlyRentAmount}}円（税別）",
        12: "{{rentPaymentDue}}",
        13: "{{dailyCalculationMethod}}",
        14: "金    {{depositAmount}}円（月額賃料    ヶ月分）",
        15: "{{specialProvisions}}",
    }
    for row_index, value in cell_values.items():
        replace_cell(terms_table.rows[row_index].cells[1], value)

    # The editable contract terms replace the static article body.  The title,
    # signature block, floor-plan heading and other legally fixed headings stay.
    paragraphs = list(document.paragraphs)
    terms_start_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.text.startswith("賃貸人 ＳＫハウジング株式会社と、賃借人")
    )
    terms_end_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.text.strip() == "第２８条（誠実協議）"
    )
    terms_start = paragraphs[terms_start_index]
    replace_paragraph(terms_start, "{{termsText}}")
    for paragraph in paragraphs[terms_start_index + 1 : terms_end_index + 2]:
        remove_paragraph(paragraph)

    # Keep the floor-plan section and populate the descriptive line.  The image
    # is a standard text tag whose merged value is HTML <img ...>, allowing the
    # server to control the snapshot without browser-supplied data.
    floor_plan_line = next(
        paragraph
        for paragraph in document.paragraphs
        if "㎡（" in paragraph.text and "坪）" in paragraph.text
    )
    replace_paragraph(
        floor_plan_line,
        "{{floorLabel}}階    {{leasedAreaSqm}}㎡（{{unitNames}}）",
    )
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "":
            previous = paragraph._element.getprevious()
            if previous is not None and "本物件平面図" in "".join(previous.itertext()):
                replace_paragraph(paragraph, "{{planImage}}")
                break

    # Editable restoration criteria replace the two fixed tables.  Keep the
    # section heading, then use a single free-flowing text tag before privacy
    # notices begin on the following page.
    restoration_heading = find_paragraph(document, "原状回復工事基準")
    restoration_body = find_paragraph(document, "■ 建築")
    replace_paragraph(restoration_body, "{{restorationText}}")
    equipment_heading = find_paragraph(document, "■ 設備")
    remove_paragraph(equipment_heading)
    for table in list(document.tables[1:]):
        remove_table(table)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def move_terms_after_contract_table(template_path: Path) -> None:
    """Separate the fixed preamble from the editable terms on the next page."""
    document = Document(template_path)
    terms_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "{{termsText}}"
    )
    replace_paragraph(
        terms_paragraph,
        "賃貸人 ＳＫハウジング株式会社と、賃借人 {{tenantName}}との間に、"
        "貸室に関する賃貸借契約（以下「本契約」という）を次のとおり締結する。",
    )

    # The first paragraph following the contract-terms table is a deliberately
    # blank spacer in the approved source. Reuse it for the editable terms and
    # force the block onto the following page.
    following = document.tables[0]._element.getnext()
    while following is not None and not following.tag.endswith('}p'):
        following = following.getnext()
    if following is None:
        raise ValueError('No paragraph was found after the contract terms table.')
    destination = Paragraph(following, document._body)
    destination.clear()
    run = destination.add_run()
    run.add_break(WD_BREAK.PAGE)
    run.add_text("{{termsText}}")
    document.save(template_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path, nargs='?')
    parser.add_argument("destination", type=Path, nargs='?')
    parser.add_argument('--move-terms-after-table', type=Path)
    args = parser.parse_args()
    if args.move_terms_after_table:
        move_terms_after_contract_table(args.move_terms_after_table)
        return
    if not args.source or not args.destination:
        parser.error('source and destination are required unless --move-terms-after-table is used')
    create_template(args.source, args.destination)


if __name__ == "__main__":
    main()
